import logging
import io
from typing import List, Optional
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from .. import models, schemas
import pypdf
import docx

logger = logging.getLogger(__name__)

class DocumentService:
    async def parse_file(self, file: UploadFile) -> str:
        content = ""
        filename = (file.filename or "").lower()
        
        try:
            file_bytes = await file.read()
            file_stream = io.BytesIO(file_bytes)
            await file.seek(0) # Reset pointer for later saving

            if filename.endswith('.pdf'):
                try:
                    reader = pypdf.PdfReader(file_stream)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            content += extracted + "\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from PDF {filename}: {e}")
            
            elif filename.endswith('.docx'):
                try:
                    doc = docx.Document(file_stream)
                    for para in doc.paragraphs:
                        content += para.text + "\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from DOCX {filename}: {e}")
            
            elif filename.endswith('.txt') or filename.endswith('.md'):
                content = file_bytes.decode('utf-8', errors='ignore')
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            return "" # Return empty string instead of failing
        finally:
            await file.seek(0)

    async def get_upload_url(
        self,
        db: AsyncSession,
        user_id: str,
        filename: str,
        content_type: str,
        title: str,
        topic_id: Optional[str] = None
    ) -> dict:
        import uuid
        from .storage_service import storage_service
        
        ext = filename.split('.')[-1] if '.' in filename else ""
        storage_key = f"{user_id}/{uuid.uuid4()}.{ext}" if ext else f"{user_id}/{uuid.uuid4()}"
        
        presigned_url = storage_service.get_presigned_upload_url(storage_key)
        
        db_document = models.Document(
            user_id=user_id,
            topic_id=topic_id,
            title=title,
            filename=filename,
            content="",
            file_type=ext.lower() if ext else "unknown",
            storage_key=storage_key,
            status="uploading"
        )
        
        db.add(db_document)
        await db.commit()
        await db.refresh(db_document)
        
        return {
            "document_id": db_document.id,
            "presigned_url": presigned_url,
            "storage_key": storage_key
        }

    async def complete_upload(
        self,
        db: AsyncSession,
        document_id: int,
        user_id: str
    ) -> models.Document:
        from sqlalchemy import update
        await db.execute(
            update(models.Document)
            .where(models.Document.id == document_id)
            .values(status="ready")
        )
        await db.commit()
        result = await db.execute(
            select(models.Document).where(models.Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        return document

    async def create_document(
        self, 
        db: AsyncSession, 
        user_id: str, 
        file: UploadFile, 
        title: str
    ) -> models.Document:
        """Save file to disk and save record to database."""
        import os
        import uuid
        
        # Create uploads directory if it doesn't exist
        upload_dir = os.path.join(os.getcwd(), "backend", "uploads")
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename to avoid collisions
        filename = file.filename or "unknown"
        ext = filename.split('.')[-1] if '.' in filename else ""
        unique_filename = f"{uuid.uuid4()}.{ext}" if ext else str(uuid.uuid4())
        file_path = os.path.join(upload_dir, unique_filename)
        
        # Save file to disk
        content_bytes = await file.read()
        with open(file_path, "wb") as f:
            f.write(content_bytes)
        
        # Seek back for parsing
        await file.seek(0)
        
        # Attempt text extraction (for search/AI) but don't fail if it's empty
        content = await self.parse_file(file)
        
        file_type = ext.lower() if ext else "unknown"
        
        db_document = models.Document(
            user_id=user_id,
            title=title,
            filename=filename,
            content=content, # Now allowed to be empty
            file_type=file_type,
            file_path=file_path
        )
        
        db.add(db_document)
        await db.commit()
        await db.refresh(db_document)
        
        return db_document

    async def get_documents(self, db: AsyncSession, user_id: str):
        from sqlalchemy import or_
        result = await db.execute(
            select(models.Document)
            .where(models.Document.user_id == user_id)
            .where(or_(models.Document.status == "ready", models.Document.status.is_(None)))
            .order_by(models.Document.created_at.desc())
        )
        return result.scalars().all()

    async def get_document(self, db: AsyncSession, document_id: int, user_id: str):
        """Get a specific document."""
        result = await db.execute(
            select(models.Document).where(models.Document.id == document_id, models.Document.user_id == user_id)
        )
        document = result.scalar_one_or_none()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        return document

    async def update_document(
        self,
        db: AsyncSession,
        document_id: int,
        user_id: str,
        update_data: schemas.DocumentUpdate
    ) -> models.Document:
        document = await self.get_document(db, document_id, user_id)
        
        data = update_data.model_dump(exclude_unset=True)
        for key, value in data.items():
            setattr(document, key, value)
            
        await db.commit()
        await db.refresh(document)
        return document

    async def delete_document(self, db: AsyncSession, document_id: int, user_id: str):
        document = await self.get_document(db, document_id, user_id)
        
        storage_key = getattr(document, "storage_key", None)
        if storage_key:
            try:
                from .storage_service import storage_service
                storage_service.delete_object(str(storage_key))
            except Exception as e:
                logger.error(f"Failed to delete object {storage_key} from MinIO: {e}")

        file_path = getattr(document, "file_path", None)
        if file_path:
            try:
                import os
                if os.path.exists(str(file_path)):
                    os.remove(str(file_path))
            except Exception as e:
                logger.error(f"Failed to delete file {file_path}: {e}")

        await db.delete(document)
        await db.commit()
        return True

document_service = DocumentService()
